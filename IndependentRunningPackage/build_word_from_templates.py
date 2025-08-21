# -*- coding: utf-8 -*-
"""
单文件版：读取本目录下三份输入与两份模板，生成与旧流程一致版式的“教案-*.docx”。

需要的文件（全部放在本脚本所在目录）：
1) 教案模板标记值-*.md         （占位映射）
2) *-18-data.json               （周次数据）
3) *-教学大纲.md               （留档，可无）
4) 教案-模板.docx               （教案头模板）
5) 课程教学教案-模板.docx       （周次表格模板，首张表作为原型）

生成：
- {科目}-教案头.docx（中间产物）
- {科目}-教师授课教案信息表集合.docx（中间产物）
- 教案-{科目}.docx（最终产物，生成在本目录）

依赖：python-docx（以及其依赖 lxml），其它仅用标准库。
"""

from __future__ import annotations
import os
import re
import sys
import json
import shutil
from copy import deepcopy
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List, Optional, Iterable

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn


# ---------- 工具与解析 ----------

def read_text(path: str, encoding: str = "utf-8") -> str:
    with open(path, "r", encoding=encoding) as f:
        return f.read()


def parse_placeholder_md(md_path: str) -> Dict[str, str]:
    """解析形如：
    - #{授课科目}：大数据基础（Hadoop）
    - #{总周数}：18
    - #{节}：1234节
    也支持：- 授课科目：大数据基础（Hadoop）
    支持中英文冒号与可选 #。
    """
    mapping: Dict[str, str] = {}
    if not os.path.exists(md_path):
        return mapping
    txt = read_text(md_path)
    # 规则1：- #{名称}：值
    line_re = re.compile(r"^\s*[-*]\s*#?\{\s*([^}\s：:]+)\s*\}\s*[：:]\s*(.+?)\s*$", re.UNICODE)
    # 规则2：允许无#的 {名称}
    line_re2 = re.compile(r"^\s*[-*]\s*\{\s*([^}\s：:]+)\s*\}\s*[：:]\s*(.+?)\s*$", re.UNICODE)
    # 规则3：允许无花括号的 名称：值
    line_re3 = re.compile(r"^\s*[-*]\s*([^{}\s：:]+)\s*[：:]\s*(.+?)\s*$", re.UNICODE)
    for raw in txt.splitlines():
        m = line_re.match(raw) or line_re2.match(raw) or line_re3.match(raw)
        if m:
            name = m.group(1).strip()
            val = m.group(2).strip()
            mapping[name] = val
    return mapping


def build_patterns_for_name(name: str) -> List[re.Pattern]:
    """支持多种占位形式（花括号/半角括号/全角括号，全角花括号，带或不带#）："""
    return [
        re.compile(r"[#＃]\s*\{\s*" + re.escape(name) + r"\s*(?:[:：][^}]+)?\s*\}", re.UNICODE),
        re.compile(r"\{\s*" + re.escape(name) + r"\s*(?:[:：][^}]+)?\s*\}", re.UNICODE),
        # 全角花括号｛｝
        re.compile(r"[#＃]?\s*｛\s*" + re.escape(name) + r"\s*(?:[:：][^｝]+)?\s*｝", re.UNICODE),
        re.compile(r"｛\s*" + re.escape(name) + r"\s*(?:[:：][^｝]+)?\s*｝", re.UNICODE),
        # 括号
        re.compile(r"[#＃]\s*\(\s*" + re.escape(name) + r"\s*(?:[:：][^)]+)?\s*\)", re.UNICODE),
        re.compile(r"\(\s*" + re.escape(name) + r"\s*(?:[:：][^)]+)?\s*\)", re.UNICODE),
        re.compile(r"[#＃]?\s*（\s*" + re.escape(name) + r"\s*(?:[:：][^）]+)?\s*）", re.UNICODE),
    ]


def _norm_label(s: str) -> str:
    s = (s or "").strip()
    s = (s.replace('\u00A0', ' ').replace('\u3000', ' ').replace('\u202F', ' ').replace('\u2007', ' ')
           .replace('\u200B', '').replace('\u200C', '').replace('\u200D', ''))
    return re.sub(r"[\s:：]", "", s)

def xml_replace_in_element(element, mapping: Dict[str, str]) -> None:
    """在给定 element 下替换所有 w:t 与 a:t 节点文本。"""
    if element is None:
        return
    W_T = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
    A_T = "{http://schemas.openxmlformats.org/drawingml/2006/main}t"
    for node in element.iter():
        if node.tag == W_T or node.tag == A_T:
            text = node.text or ""
            new_text = text
            for key, val in mapping.items():
                for pat in build_patterns_for_name(key):
                    new_text = pat.sub(str(val), new_text)
            if new_text != text:
                node.text = new_text


def xml_replace_in_doc(doc: Document, mapping: Dict[str, str]) -> None:
    xml_replace_in_element(doc.element.body, mapping)
    for sect in doc.sections:
        if sect.header and getattr(sect.header, "_element", None) is not None:
            xml_replace_in_element(sect.header._element, mapping)
        if sect.footer and getattr(sect.footer, "_element", None) is not None:
            xml_replace_in_element(sect.footer._element, mapping)


def _norm_label(s: str) -> str:
    return re.sub(r"[\s:：]", "", (s or "").strip())


def write_cell_text_preserve_style(cell, val: str) -> None:
    """将整格内容替换为 val，并尽量保留该格首个 run 的字体/样式。"""
    paras = list(cell.paragraphs)
    if not paras:
        p = cell.add_paragraph("")
        r = p.add_run("")
    else:
        p = paras[0]
        runs = list(p.runs)
        if runs:
            r = runs[0]
            for extra in runs[1:]:
                p._p.remove(extra._r)
        else:
            r = p.add_run("")
        for extra_p in paras[1:]:
            cell._tc.remove(extra_p._p)
    r.text = val if val is not None else ""


def replace_placeholders_in_all_cells(doc: Document, mapping: Dict[str, str]) -> None:
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                original = cell.text or ""
                new_text = original
                for key, val in mapping.items():
                    for pat in build_patterns_for_name(key):
                        new_text = pat.sub(str(val), new_text)
                if new_text != original:
                    write_cell_text_preserve_style(cell, new_text)


def fill_tables_by_labels(doc: Document, mapping: Dict[str, str]) -> None:
    keys = ["授课科目", "授课老师", "授课班级", "授课起止时间", "周学时", "考核方式"]
    norm_keys = {_norm_label(k): k for k in keys}
    for tbl in doc.tables:
        for row in tbl.rows:
            if len(row.cells) < 2:
                continue
            label_norm = _norm_label(row.cells[0].text)
            if label_norm in norm_keys:
                k = norm_keys[label_norm]
                v = str(mapping.get(k, "") or "")
                write_cell_text_preserve_style(row.cells[1], v)


def limit_text(s: str, max_len: int) -> str:
    s = (s or "").strip()
    if max_len <= 0:
        return s
    return s[:max_len]


def cleanup_midline_spaces(s: str) -> str:
    # 将多行压缩为单行，并折叠连续空格
    s = re.sub(r"\n+", " ", s or "")
    s = re.sub(r" {2,}", " ", s)
    return s.strip()


def ensure_week_word_in_time(s: str) -> str:
    txt = (s or '').strip()
    if not txt:
        return txt
    txt_norm = (
        txt.replace('\u00A0', ' ').replace('\u3000', ' ').replace('\u202F', ' ').replace('\u2007', ' ')
           .replace('\u200B', '').replace('\u200C', '').replace('\u200D', '')
    )
    txt_norm = re.sub(r"\s+", " ", txt_norm).strip()
    if '周' in txt_norm:
        return txt_norm
    m = re.match(r'^第\s*(\d+)\s+([0-9]+)\s*节$', txt_norm)
    if m:
        week = m.group(1)
        sections = m.group(2) + '节'
        return f"第 {week} 周 {sections}"
    return txt_norm


def fix_time_cell_for_table(tbl) -> None:
    for row in tbl.rows:
        cells = list(row.cells)
        for i, c in enumerate(cells):
            if _norm_label(c.text) == _norm_label("授课时间") and i + 1 < len(cells):
                data_cell = cells[i+1]
                original = data_cell.text or ""
                new_text = ensure_week_word_in_time(cleanup_midline_spaces(original))
                if new_text != original.strip():
                    write_cell_text_preserve_style(data_cell, new_text)


def derive_week_hours(section_value: str) -> str:
    v = (section_value or "").strip().replace(" ", "")
    four_set = {"1234节", "5678节"}
    return "4" if v in four_set else "2"


def parse_font_size_pt(s: str | None) -> float | None:
    if not s:
        return None
    txt = str(s).strip().lower()
    m = re.match(r"^([0-9]+(?:\.[0-9]+)?)\s*pt$", txt)
    if m:
        try:
            return float(m.group(1))
        except Exception:
            return None
    m2 = re.match(r"^([0-9]+(?:\.[0-9]+)?)$", txt)
    if m2:
        try:
            return float(m2.group(1))
        except Exception:
            return None
    zh_map = {"三号": 16.0, "小三": 15.0, "四号": 14.0, "小四": 12.0, "五号": 10.5, "小五": 9.0}
    return zh_map.get(str(s).strip(), None)


def unify_document_font(doc: Document, font_name: Optional[str] = None, font_size_pt: Optional[float] = None) -> None:
    # 选择字体名
    if not font_name:
        for p in doc.paragraphs:
            for r in p.runs:
                if r.font and r.font.name:
                    font_name = r.font.name
                    break
            if font_name:
                break
        if not font_name:
            font_name = "宋体"

    def _set_run_font(run):
        try:
            if font_name:
                run.font.name = font_name
                r = run._element
                rPr = r.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:eastAsia'), font_name)
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
            if font_size_pt:
                run.font.size = Pt(font_size_pt)
        except Exception:
            pass

    for p in doc.paragraphs:
        for run in p.runs:
            _set_run_font(run)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        _set_run_font(run)


def unify_document_font_excluding(doc: Document, font_name: Optional[str], font_size_pt: Optional[float],
                                  exclude_texts: Iterable[str]) -> None:
    def _norm_text(s: str) -> str:
        if s is None:
            return ""
        s = s.replace('\u00A0', '').replace('\u3000', '')
        s = ''.join(ch for ch in s if not ch.isspace())
        return s.strip().lower()

    exclude_norm = {_norm_text(t) for t in exclude_texts if t}

    def _should_exclude(paragraph) -> bool:
        txt = ''.join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
        return _norm_text(txt) in exclude_norm

    def _set_run_font(run):
        try:
            if font_name:
                run.font.name = font_name
                r = run._element
                rPr = r.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(qn('w:eastAsia'), font_name)
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
            if font_size_pt:
                run.font.size = Pt(font_size_pt)
        except Exception:
            pass

    for p in doc.paragraphs:
        if _should_exclude(p):
            continue
        for run in p.runs:
            _set_run_font(run)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if _should_exclude(p):
                        continue
                    for run in p.runs:
                        _set_run_font(run)


# ---------- 主流程（单文件实现） ----------

def find_input_files(src_dir: Path) -> tuple[Path, Path, Path | None]:
    """在独立运行包中优先从 data/ 目录寻找 MD 与 JSON；若未找到，再回落到脚本同级目录。"""
    search_dirs = [
        (src_dir / "data").resolve(),  # 独立运行包的数据目录
        src_dir.resolve(),               # 脚本所在目录（向后兼容）
    ]

    # 先尝试精确示例名（若存在则优先）
    md_path = json_path = syllabus_path = None
    for base in search_dirs:
        t_md = base / "教案模板标记值-大数据基础（Hadoop）.md"
        t_json = base / "大数据基础（Hadoop）-18-data.json"
        t_syllabus = base / "大数据基础（Hadoop）-教学大纲.md"
        if md_path is None and t_md.exists():
            md_path = t_md
        if json_path is None and t_json.exists():
            json_path = t_json
        if syllabus_path is None and t_syllabus.exists():
            syllabus_path = t_syllabus

    # 若未命中精确示例名，则使用通配优先匹配
    if md_path is None:
        for base in search_dirs:
            md_candidates = list(base.glob("教案模板标记值-*.md"))
            if md_candidates:
                md_path = md_candidates[0]
                break

    if json_path is None:
        for base in search_dirs:
            json_candidates = list(base.glob("*-*-data.json")) + list(base.glob("*-data.json"))
            if json_candidates:
                json_path = json_candidates[0]
                break

    if syllabus_path is None:
        for base in search_dirs:
            syllabus_candidates = list(base.glob("*-教学大纲.md"))
            if syllabus_candidates:
                syllabus_path = syllabus_candidates[0]
                break

    if not md_path or not md_path.exists():
        raise FileNotFoundError(
            "未找到标记值MD：请将 ‘教案模板标记值-*.md’ 放到 data/ 目录或脚本同级目录"
        )
    if not json_path or not json_path.exists():
        raise FileNotFoundError(
            "未找到周次JSON：请将 ‘*-data.json’ 放到 data/ 目录或脚本同级目录"
        )
    if syllabus_path and (not syllabus_path.exists()):
        syllabus_path = None
    return md_path, json_path, syllabus_path


def find_docx_templates(src_dir: Path) -> tuple[Path, Path]:
    """优先在独立运行包的 templates/ 目录下寻找两份 docx 模板；若不存在，尝试在常见相对路径下查找。"""
    candidates = [
        src_dir.resolve(),
        (src_dir / "templates").resolve(),            # IndependentRunningPackage/templates
        (src_dir / "templates" / "教案").resolve(),   # 允许子目录教案
        # UI/static/templates -> UI/templates/教案（部分仓库可能放这里）
        (src_dir / ".." / ".." / "templates" / "教案").resolve(),
        # UI/static/templates -> Generator/templates/教案（本仓库实际位置）
        (src_dir / ".." / ".." / ".." / "templates" / "教案").resolve(),
    ]
    for base in candidates:
        p1 = base / "教案-模板.docx"
        p2 = base / "课程教学教案-模板.docx"
        if p1.exists() and p2.exists():
            return p1, p2
    raise FileNotFoundError(
        "缺少模板：请将 ‘教案-模板.docx’ 与 ‘课程教学教案-模板.docx’ 放到脚本同级目录，或放入 templates/、templates/教案、UI/templates/教案 或 Generator/templates/教案 目录。"
    )


def build_head_doc(head_tpl: Path, mapping: Dict[str, str], out_dir: Path, subject: str) -> Path:
    doc = Document(str(head_tpl))
    # 默认补全
    try:
        total_weeks = int(str(mapping.get("总周数", "16")).strip())
    except Exception:
        total_weeks = 16
    if not mapping.get("授课起止时间"):
        mapping["授课起止时间"] = f"第1周-第{total_weeks}周"
    if not mapping.get("周学时"):
        mapping["周学时"] = derive_week_hours(mapping.get("节", "1234节"))
    if not mapping.get("考核方式"):
        mapping["考核方式"] = "考察"

    # 全局替换与兜底表格填充
    xml_replace_in_doc(doc, mapping)
    replace_placeholders_in_all_cells(doc, mapping)
    fill_tables_by_labels(doc, mapping)

    out = out_dir / f"{subject}-教案头.docx"
    out_dir.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def get_time_cell_font_from_table(tbl) -> tuple[str | None, float | None]:
    font_name, font_size = None, None
    try:
        for row in tbl.rows:
            cells = list(row.cells)
            for i, c in enumerate(cells):
                if _norm_label(c.text) == _norm_label("授课时间") and i + 1 < len(cells):
                    data_cell = cells[i+1]
                    for p in data_cell.paragraphs:
                        for r in p.runs:
                            if not font_name and r.font and r.font.name:
                                font_name = r.font.name
                            if not font_size and r.font and r.font.size:
                                try:
                                    font_size = float(r.font.size.pt)
                                except Exception:
                                    pass
                            if font_name and font_size:
                                return font_name, font_size
        return font_name, font_size
    except Exception:
        return None, None


def append_table_from_template(doc: Document, tpl_table):
    new_tbl = deepcopy(tpl_table._tbl)
    doc._body._element.append(new_tbl)
    return doc.tables[-1]


def replace_placeholders_in_table_cells(tbl, mapping: Dict[str, str]) -> None:
    for row in tbl.rows:
        for cell in row.cells:
            original = cell.text or ""
            new_text = original
            for key, val in mapping.items():
                for pat in build_patterns_for_name(key):
                    new_text = pat.sub(str(val), new_text)
            if new_text != original:
                write_cell_text_preserve_style(cell, new_text)


def build_weeks_doc(week_tpl: Path, mapping: Dict[str, str], json_path: Path, out_dir: Path, subject: str) -> Path:
    if not week_tpl.exists():
        raise FileNotFoundError(f"未找到周表格模板: {week_tpl}")

    # 加载周数据
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    try:
        total_weeks = int(str(mapping.get("总周数", data.get("总周数", "16"))).strip())
    except Exception:
        total_weeks = int(data.get("总周数", 16))
    sections = str(mapping.get("节", data.get("节", "1234节"))).strip()

    # 默认补全
    if not mapping.get("授课起止时间"):
        mapping["授课起止时间"] = f"第1周-第{total_weeks}周"
    if not mapping.get("周学时"):
        mapping["周学时"] = derive_week_hours(sections)
    if not mapping.get("考核方式"):
        mapping["考核方式"] = "平时30%+期末(或大作业)70%"

    # 原型表格来自周模板
    week_tpl_doc = Document(str(week_tpl))
    if not week_tpl_doc.tables:
        raise RuntimeError("周表格模板文档中未找到表格")
    week_table_tpl = week_tpl_doc.tables[0]

    # 从用户/模板确定目标字体与字号
    user_font_name = mapping.get("统一字体名称", "").strip() or None
    user_font_size_pt = parse_font_size_pt(mapping.get("统一字号", None))
    tpl_font_name, tpl_font_pt = get_time_cell_font_from_table(week_table_tpl)
    chosen_font_name = user_font_name or tpl_font_name or "宋体"
    chosen_font_pt = user_font_size_pt or tpl_font_pt

    # 以周模板文档作为基底并清空正文
    base_doc = Document(str(week_tpl))
    body = base_doc._body._element
    for child in list(body):
        body.remove(child)

    # 若周模板包含页眉/页脚占位，先全局替换
    xml_replace_in_doc(base_doc, mapping)

    # 准备周数组
    weeks: List[Dict[str, Any]] = list(data.get("周次", []))
    if len(weeks) < total_weeks:
        weeks = weeks + [{} for _ in range(total_weeks - len(weeks))]
    else:
        weeks = weeks[:total_weeks]

    for idx, wk in enumerate(weeks, start=1):
        wk_mapping: Dict[str, str] = {}
        # 复制基础字段
        for key in ["授课科目", "授课老师", "授课班级", "人数", "授课起止时间", "考核方式", "授课地点"]:
            if key in mapping:
                wk_mapping[key] = mapping[key]
        # 别名：班级人数 -> 人数
        if "人数" not in wk_mapping:
            alias_people = mapping.get("人数") or mapping.get("班级人数")
            if alias_people:
                wk_mapping["人数"] = alias_people
        # 每周字段
        wk_mapping["单元"] = f"{idx} 单元"
        wk_mapping["周"] = f"{idx}"
        wk_mapping["节"] = sections
        wk_mapping["授课时间"] = f"第 {idx} 周"
        wk_mapping["课题"] = limit_text(str(wk.get("课题", "")), 50)
        wk_mapping["教学目标"] = str(wk.get("教学目标", ""))
        wk_mapping["教学重点"] = str(wk.get("教学重点", ""))
        wk_mapping["教学难点"] = str(wk.get("教学难点", ""))
        wk_mapping["授课内容1"] = str(wk.get("授课内容1", ""))
        wk_mapping["授课内容2"] = str(wk.get("授课内容2", ""))
        wk_mapping["授课内容3"] = str(wk.get("授课内容3", ""))
        wk_mapping["授课内容4"] = str(wk.get("授课内容4", ""))
        wk_mapping["课后小结"] = ""
        wk_mapping["作业"] = str(wk.get("作业", ""))

        # 插入一份周表格
        new_tbl = append_table_from_template(base_doc, week_table_tpl)
        # 先尝试 XML 级替换
        xml_replace_in_element(new_tbl._tbl, wk_mapping)
        xml_replace_in_element(new_tbl._tbl, mapping)
        # 再做逐表格 cell 级兜底替换（处理占位符被拆分到多 w:t 的情况）
        merged_map = dict(mapping)
        merged_map.update(wk_mapping)
        replace_placeholders_in_table_cells(new_tbl, merged_map)
        # 兜底修正
        fix_time_cell_for_table(new_tbl)

        # 分页
        if idx < len(weeks):
            p = base_doc.add_paragraph("")
            run = p.add_run("")
            run.add_break(WD_BREAK.PAGE)

    # 统一字体（按用户/模板选择）
    unify_document_font(base_doc, chosen_font_name, chosen_font_pt)

    # 保存
    out = out_dir / f"{subject}-教师授课教案信息表集合.docx"
    out_dir.mkdir(parents=True, exist_ok=True)
    try:
        base_doc.save(str(out))
    except PermissionError:
        ts = datetime.now().strftime("%Y%m%d-%H%M%S")
        out = out_dir / f"{subject}-教师授课教案信息表集合-{ts}.docx"
        base_doc.save(str(out))
    return out


def merge_docs(head_doc_path: Path, append_doc_path: Path, out_path: Path, font_name: Optional[str], font_size_pt: Optional[float]) -> None:
    head_doc = Document(str(head_doc_path))
    append_doc = Document(str(append_doc_path))

    for element in append_doc.element.body:
        head_doc.element.body.append(deepcopy(element))

    exclude_titles = [
        "广 州 现 代 信 息 工 程 职 业 技 术 学 院",
        "广州现代信息工程职业技术学院",
        "教 师 授 课 教 案",
        "教师授课教案",
        "教师授课教案信息表",
    ]
    unify_document_font_excluding(head_doc, font_name, font_size_pt, exclude_titles)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    head_doc.save(str(out_path))


def main() -> None:
    src_dir = Path(__file__).parent.resolve()

    # 1) 输入与模板
    md_path, json_path, _syllabus = find_input_files(src_dir)
    head_tpl, week_tpl = find_docx_templates(src_dir)

    # 2) 映射与关键字段
    base_mapping = parse_placeholder_md(str(md_path))
    subject = (base_mapping.get("授课科目") or "").strip()
    if not subject:
        # 若 MD 未给，尝试从 JSON 内容推断
        with open(json_path, 'r', encoding='utf-8') as f:
            jd = json.load(f)
        subject = (jd.get("授课科目") or "").strip()
    if not subject:
        raise RuntimeError("未在标记值MD或JSON中找到 ‘授课科目’")

    # 3) 生成教案头
    tmp_dir = src_dir / "_tmp_build"
    head_doc = build_head_doc(head_tpl, dict(base_mapping), tmp_dir, subject)

    # 4) 生成周次集合
    weeks_doc = build_weeks_doc(week_tpl, dict(base_mapping), json_path, tmp_dir, subject)

    # 5) 合并并统一字体（可从映射读取用户配置）
    user_font_name = (base_mapping.get("统一字体名称") or "").strip() or None
    user_font_size_pt = parse_font_size_pt(base_mapping.get("统一字号"))
    final_doc = src_dir / f"教案-{subject}.docx"
    merge_docs(head_doc, weeks_doc, final_doc, user_font_name, user_font_size_pt)

    # 6) 清理临时产物
    try:
        shutil.rmtree(tmp_dir)
    except Exception:
        pass

    print(f"[完成] 生成成功：{final_doc}")


if __name__ == "__main__":
    main()